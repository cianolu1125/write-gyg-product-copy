#!/usr/bin/env python3
import argparse
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from validate_gyg_copy import HARD, validate

BLUE = RGBColor(25, 78, 121)
LIGHT_BLUE = "DDEBF7"
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p

def add_bullets(doc, items):
    if not items:
        doc.add_paragraph("Not applicable.")
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")

def build(data, output):
    errors, warnings, counts = validate(data)
    if errors:
        raise SystemExit("Validation failed:\n- " + "\n- ".join(errors))

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 16), ("Heading 2", 12)):
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = BLUE
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(10)
        styles[name].paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GetYourGuide Product Copy")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE
    sub = doc.add_paragraph(data["product_title"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.bold = True
    sub.paragraph_format.space_after = Pt(14)

    add_heading(doc, "1. Product Title")
    doc.add_paragraph(data["product_title"])
    add_heading(doc, "2. Short Description")
    doc.add_paragraph(data["short_description"])
    add_heading(doc, "3. Highlights")
    add_bullets(doc, data["highlights"])

    add_heading(doc, "4. Full Description")
    full = data["full_description"]
    stops = full.get("stops", [])
    for index, stop in enumerate(stops):
        h = doc.add_paragraph()
        h.paragraph_format.keep_with_next = True
        r = h.add_run(stop["heading"])
        r.bold = True
        r.font.color.rgb = BLUE
        for detail in stop.get("details", []):
            p = doc.add_paragraph(str(detail), style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
        doc.add_paragraph(stop["body"])
        if index < len(stops) - 1:
            doc.add_paragraph()

    sections = [
        ("5. Includes", "includes"),
        ("6. Not Included", "not_includes"),
        ("7. What to Bring", "what_to_bring"),
        ("8. Not Allowed", "not_allowed"),
        ("9. Know Before You Go", "know_before_you_go"),
        ("10. Voucher Information", "voucher_information"),
    ]
    for heading, key in sections:
        add_heading(doc, heading)
        add_bullets(doc, data.get(key, []))

    add_heading(doc, "11. Booking Options")
    options = data.get("options", [])
    if not options:
        doc.add_paragraph("No separate booking options supplied.")
    for option in options:
        add_heading(doc, option.get("title", "Option"), level=2)
        if option.get("description"):
            doc.add_paragraph(option["description"])
        if option.get("includes"):
            p = doc.add_paragraph()
            p.add_run("Option Includes:").bold = True
            for item in option["includes"]:
                doc.add_paragraph(str(item), style="List Bullet")
        if option.get("not_includes"):
            p = doc.add_paragraph()
            p.add_run("Option Does Not Include:").bold = True
            for item in option["not_includes"]:
                doc.add_paragraph(str(item), style="List Bullet")
        if option.get("meeting_pickup"):
            p = doc.add_paragraph()
            p.add_run("Meeting/Pickup: ").bold = True
            p.add_run(option["meeting_pickup"])
        if option.get("availability"):
            p = doc.add_paragraph()
            p.add_run("Availability: ").bold = True
            p.add_run(option["availability"])

    add_heading(doc, "12. Validation Summary")
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(2.6)
    table.columns[1].width = Inches(1.3)
    table.columns[2].width = Inches(2.3)
    headers = ("Field", "Characters", "Rule")
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
    rows = [
        ("Product Title", counts["product_title"], f"{HARD['product_title']['min']}-{HARD['product_title']['max']}"),
        (
            "Short Description",
            counts["short_description"],
            f"{HARD['short_description']['min']}-{HARD['short_description']['max']}; 2-3 sentences",
        ),
        (
            "Full Description",
            counts["full_description"],
            f"{HARD['full_description']['min']:,}-{HARD['full_description']['max']:,}",
        ),
    ]
    rows += [
        (f"Highlight {i}", n, f"{HARD['highlights']['min_each']}-{HARD['highlights']['max_each']}")
        for i, n in enumerate(counts["highlights"], 1)
    ]
    rows += [
        (f"Stop {i} body", n, f"{HARD['stop_body']['min']}-{HARD['stop_body']['max']}")
        for i, n in enumerate(counts["stop_bodies"], 1)
    ]
    for field, n, rule in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = str(field), str(n), rule
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    status = doc.add_paragraph()
    status.add_run("Validation status: PASS").bold = True
    if warnings:
        status.add_run(". Review warnings: " + "; ".join(warnings))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

def main():
    parser = argparse.ArgumentParser(description="Build a polished Word document from validated GYG copy JSON.")
    parser.add_argument("json_file", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    data = json.loads(args.json_file.read_text(encoding="utf-8-sig"))
    build(data, args.output_docx)
    print(args.output_docx.resolve())

if __name__ == "__main__":
    main()
